import os
import uuid
from typing import List, Dict, Any
from datetime import datetime

import PyPDF2
import docx
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer

class Document:
    """Class to represent a document with metadata and content."""
    def __init__(self, id: str, filename: str, content: str, metadata: Dict = None):
        self.id = id
        self.filename = filename
        self.content = content
        self.metadata = metadata or {}
        self.chunks = []
    
    def add_chunk(self, chunk):
        """Add a chunk to the document."""
        self.chunks.append(chunk)
        chunk.document_id = self.id
        
    def to_dict(self):
        """Convert document to dictionary."""
        return {
            "id": self.id,
            "filename": self.filename,
            "metadata": self.metadata,
            "chunks": len(self.chunks)
        }

class Chunk:
    """Class to represent a chunk of text from a document."""
    def __init__(self, id: str, text: str, metadata: Dict = None):
        self.id = id
        self.text = text
        self.metadata = metadata or {}
        self.document_id = None  # Set when added to a document

class DocumentProcessor:
    """Process documents to extract text and create chunks."""
    
    def __init__(self, upload_folder: str = "uploads"):
        """Initialize with upload folder path."""
        self.upload_folder = upload_folder
        os.makedirs(upload_folder, exist_ok=True)
        
    def extract_text(self, file_path: str) -> str:
        """Extract text from a file based on its extension."""
        _, file_extension = os.path.splitext(file_path)
        
        if file_extension.lower() == '.pdf':
            return self._extract_text_from_pdf(file_path)
        elif file_extension.lower() == '.docx':
            return self._extract_text_from_docx(file_path)
        elif file_extension.lower() == '.txt':
            return self._extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file."""
        text = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text() + "\n"
        return text
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from a TXT file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return text
    
    def process_document(self, file_path: str, filename: str) -> Document:
        """Process a document to extract text and create document object."""
        content = self.extract_text(file_path)
        
        # Create a document with unique ID
        document_id = str(uuid.uuid4())
        document = Document(
            id=document_id,
            filename=filename,
            content=content,
            metadata={"uploaded_at": datetime.now().isoformat()}
        )
        
        # Create chunks from the document
        self.chunk_document(document)
        
        return document
    
    def chunk_document(self, document: Document, chunk_size: int = 1000, overlap: int = 200) -> List[Chunk]:
        """Split the document into overlapping chunks."""
        content = document.content
        chunks = []
        
        # Simple chunking by characters with overlap
        for i in range(0, len(content), chunk_size - overlap):
            chunk_text = content[i:i + chunk_size]
            if len(chunk_text) < 50:  # Skip very small chunks
                continue
            
            chunk_id = f"{document.id}_{len(chunks)}"
            chunk = Chunk(
                id=chunk_id,
                text=chunk_text,
                metadata={
                    "position": len(chunks),
                    "start_char": i,
                    "end_char": i + len(chunk_text)
                }
            )
            document.add_chunk(chunk)
            chunks.append(chunk)
        
        return chunks
    
    def generate_embeddings(self, chunks: List[Chunk]) -> Dict[str, np.ndarray]:
        """Generate TF-IDF embeddings for document chunks."""
        # Extract text from chunks
        texts = [chunk.text for chunk in chunks]
        
        # Generate TF-IDF vectors
        vectorizer = TfidfVectorizer(max_features=300)
        tfidf_matrix = vectorizer.fit_transform(texts)
        
        # Store embeddings for each chunk
        embeddings = {}
        for i, chunk in enumerate(chunks):
            vector = tfidf_matrix[i].toarray()[0]
            # Normalize vector
            norm = np.linalg.norm(vector)
            if norm > 0:
                vector = vector / norm
            embeddings[chunk.id] = vector
        
        return embeddings, vectorizer